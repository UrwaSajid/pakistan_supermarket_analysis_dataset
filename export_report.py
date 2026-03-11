"""
export_report.py — Convert REPORT.md to a styled Word document (REPORT.docx)

Usage:
    python export_report.py

Output:
    REPORT.docx in the project root directory
"""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Colour palette (matches the Streamlit dashboard theme)
# ---------------------------------------------------------------------------
COLOUR_H1   = RGBColor(0x0f, 0x17, 0x2a)   # dark navy  #0f172a
COLOUR_H2   = RGBColor(0x1d, 0x4e, 0xd8)   # strong blue #1d4ed8
COLOUR_H3   = RGBColor(0x1e, 0x40, 0xaf)   # mid blue   #1e40af
COLOUR_CODE_BG = "F1F5F9"                   # light slate for code blocks
COLOUR_TH_BG   = "1D4ED8"                   # blue table header
COLOUR_TR_ALT  = "EFF6FF"                   # alternating row tint


# ---------------------------------------------------------------------------
# Helpers — low-level XML manipulation
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Fill a table cell background with a solid hex colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(table):
    """Add thin borders to every cell in a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFDBFE")   # light blue border
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _apply_shading_paragraph(para, hex_color: str):
    """Give a paragraph a background shading (used for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


# ---------------------------------------------------------------------------
# Inline markdown  →  docx runs  (bold / italic / inline-code)
# ---------------------------------------------------------------------------

def _add_inline(para, text: str, base_size: int = 11, base_font: str = "Calibri"):
    """
    Parse a string with markdown bold (**text**), italic (*text*),
    and inline-code (`text`) and add styled runs to *para*.
    """
    # Combined regex: code first so ` inside ** ** is handled correctly
    pattern = re.compile(
        r"(`[^`]+`)"          # inline code
        r"|(\*\*[^*]+\*\*)"  # bold
        r"|(\*[^*]+\*)"       # italic
        r"|([^`*]+)"          # plain text
    )
    for m in pattern.finditer(text):
        code_tok, bold_tok, ital_tok, plain_tok = m.groups()
        run = para.add_run()
        run.font.name = base_font
        run.font.size = Pt(base_size)

        if code_tok:
            run.text = code_tok[1:-1]          # strip backticks
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)  # red
        elif bold_tok:
            run.text = bold_tok[2:-2]          # strip **
            run.bold = True
        elif ital_tok:
            run.text = ital_tok[1:-1]          # strip *
            run.italic = True
        else:
            run.text = plain_tok or ""


# ---------------------------------------------------------------------------
# REPORT.md  →  Document builder
# ---------------------------------------------------------------------------

def build_document(md_path: Path) -> Document:
    doc = Document()

    # ---- page margins (narrower than default for readability) ----
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ---- default body style ----
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    total = len(lines)
    i = 0

    while i < total:
        line = lines[i]

        # ----------------------------------------------------------------
        # H1  →  Title paragraph
        # ----------------------------------------------------------------
        if line.startswith("# ") and not line.startswith("## "):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line[2:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = COLOUR_H1
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.space_before = Pt(18)
            i += 1
            continue

        # ----------------------------------------------------------------
        # H2  →  Heading 1 style
        # ----------------------------------------------------------------
        if line.startswith("## ") and not line.startswith("### "):
            para = doc.add_paragraph()
            run = para.add_run(line[3:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = COLOUR_H2
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after = Pt(6)
            # bottom border for H2
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1D4ED8")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ----------------------------------------------------------------
        # H3  →  Heading 2 style
        # ----------------------------------------------------------------
        if line.startswith("### "):
            para = doc.add_paragraph()
            run = para.add_run(line[4:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = COLOUR_H3
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ----------------------------------------------------------------
        # H4  →  small bold label
        # ----------------------------------------------------------------
        if line.startswith("#### "):
            para = doc.add_paragraph()
            run = para.add_run(line[5:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = COLOUR_H3
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ----------------------------------------------------------------
        # Horizontal rule  ---
        # ----------------------------------------------------------------
        if re.match(r"^---+\s*$", line):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CBD5E1")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ----------------------------------------------------------------
        # Code block  ``` ... ```
        # ----------------------------------------------------------------
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < total and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            for cl in code_lines:
                para = doc.add_paragraph()
                _apply_shading_paragraph(para, COLOUR_CODE_BG)
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                run = para.add_run(cl)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            # spacer after block
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # ----------------------------------------------------------------
        # Markdown table  |---|---|
        # ----------------------------------------------------------------
        if line.strip().startswith("|"):
            # Collect all consecutive table lines
            tbl_lines = []
            while i < total and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1

            # Filter out the separator row (|---|---|)
            data_rows = [r for r in tbl_lines if not re.match(r"^\s*\|[\s\-|:]+\|\s*$", r)]
            if not data_rows:
                continue

            # Parse rows into cells
            parsed = []
            for row in data_rows:
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                parsed.append(cells)

            if not parsed:
                continue

            n_cols = max(len(r) for r in parsed)
            # Normalise all rows to same column count
            parsed = [r + [""] * (n_cols - len(r)) for r in parsed]

            tbl = doc.add_table(rows=len(parsed), cols=n_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_cell_borders(tbl)

            for ri, row_data in enumerate(parsed):
                for ci, cell_text in enumerate(row_data):
                    cell = tbl.cell(ri, ci)
                    cell.paragraphs[0].clear()
                    para = cell.paragraphs[0]

                    if ri == 0:
                        # Header row — white bold text on blue
                        _set_cell_bg(cell, COLOUR_TH_BG)
                        run = para.add_run(re.sub(r"\*+", "", cell_text))
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    else:
                        if ri % 2 == 0:
                            _set_cell_bg(cell, COLOUR_TR_ALT)
                        _add_inline(para, cell_text, base_size=10)

                    para.paragraph_format.space_after = Pt(2)
                    para.paragraph_format.space_before = Pt(2)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # ----------------------------------------------------------------
        # Bullet list  - item  or  * item
        # ----------------------------------------------------------------
        if re.match(r"^(\s*[-*])\s+", line):
            indent = len(line) - len(line.lstrip())
            content = re.sub(r"^\s*[-*]\s+", "", line)
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25 + indent * 0.1)
            para.paragraph_format.space_after = Pt(2)
            para.clear()
            _add_inline(para, content)
            i += 1
            continue

        # ----------------------------------------------------------------
        # Numbered list  1. item
        # ----------------------------------------------------------------
        if re.match(r"^\d+\.\s+", line):
            content = re.sub(r"^\d+\.\s+", "", line)
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.space_after = Pt(2)
            para.clear()
            _add_inline(para, content)
            i += 1
            continue

        # ----------------------------------------------------------------
        # Blank line
        # ----------------------------------------------------------------
        if not line.strip():
            # Add a small spacer only if the previous paragraph wasn't already a spacer
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ----------------------------------------------------------------
        # Normal paragraph (possibly containing inline markdown)
        # ----------------------------------------------------------------
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        _add_inline(para, line.strip())
        i += 1

    return doc


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    here = Path(__file__).parent
    md_file   = here / "REPORT.md"
    docx_file = here / "REPORT.docx"

    if not md_file.exists():
        raise FileNotFoundError(f"Cannot find {md_file}")

    print(f"Reading  {md_file} …")
    doc = build_document(md_file)

    print(f"Saving   {docx_file} …")
    doc.save(docx_file)

    size_kb = docx_file.stat().st_size / 1024
    print(f"Done!    REPORT.docx created ({size_kb:.1f} KB)")
    print(f"Path:    {docx_file.resolve()}")


if __name__ == "__main__":
    main()
